# Import statement
from docx import Document
import pandas as pd
import csv
from os import path


# ======================================================================================================================
# Helper functions
# ======================================================================================================================
# Function to save a table as CSV using the csv library
def save_table_as_csv(table, file_path):
    with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            writer.writerow(row_data)


# Convert the table to a Pandas DF
def save_table_info(table):
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        data.append(row_data)
    df = pd.DataFrame(data)
    return df


# Alternate function to save a table as CSV using pandas
def save_tables_to_csv(table, output_filename):
    table.to_csv(output_filename, index=True)


# ======================================================================================================================
# Main Function
# ======================================================================================================================
def extract_tables(document, pandas_tables=False, output_location=''):
    """
    This tool reads an ADD word document and extracts the tables in the Appendix section.
    :param document: path of the Word document; STRING
    :param pandas_tables: default False. If false the output files will be created. If True, output files will not be
    created. A dictionary of pandas dataframes representing the table(s) will be returned with the key being the
    appendix name; BOOL
    :param output_location: Output file location for the table(s). Do not specify the name, will e named the Appendix
    title; STRING
    :return: (conditionally) A dictionary of pandas dataframes representing the table(s) in the appendix section with
     the key being the name of the appendix
    """

    # Convert the Word document into a docx object
    doc = Document(document)

    # List of appendix titles to look for
    appendices = ['Appendix A', 'Appendix B', 'Appendix C', 'Appendix D', 'Appendix E', 'Appendix F', 'Appendix G',
                  'Appendix H', 'Appendix I', 'Appendix J', 'Appendix K', 'Appendix L', 'Appendix M', 'Appendix N',
                  'Appendix O', 'Appendix P', 'Appendix Q', 'Appendix R', 'Appendix S', 'Appendix T', 'Appendix U',
                  'Appendix V', 'Appendix W', 'Appendix X', 'Appendix Y', 'Appendix Z']

    # Reset the index to 0
    index_counter = 0
    # Black list that will store the index of the table location in the Word document
    tbl_index_list = []

    # Loop through the document and look at the blocks of information
    for block in doc.element.body:
        # Determine which blocks are tables
        if block.tag.endswith('tbl'):
            # Mark the index of the blocks that are tables
            tbl_index_list.append(index_counter)
        # Move the counter by 1
        index_counter += 1

    # Make a list for the index of the text paragraphs that comes directly before the tables
    p_index_list = [num - 1 for num in tbl_index_list]

    # Reset the index to 0
    index_counter = 0

    # Empty list that will be filled with smaller lists that represent the index of the block for the table, appendix
    # name, and eventually the index of the table in the list of tables in the document
    title_list = []
    for block in doc.element.body:
        # Mini list that will be appended to the title_list that contain the important info
        temp_list = []
        # Go through the P index list
        if index_counter in p_index_list:
            # get the text for the paragraph
            para_text = block.text.strip()
            # Get the first two words to ensure it is an Appendix (could be an appendix elsewhere though)
            words = para_text.split()
            beginning = " ".join(words[:2])

            # Checks if the first two words are 'Appendix <Letter>'
            if beginning in appendices:
                # Make the appendix title as a file name
                appendix_title = para_text.replace(" ", "_")
                tbl_index = index_counter + 1  # table block is one after the Appendix block
                # populate the mini list
                temp_list.append(tbl_index)
                temp_list.append(appendix_title)
                # add the mini list to the title_list
                title_list.append(temp_list)

        # Increase the index by one
        index_counter += 1

    # Append the value of the index of the table for the list of tables
    for place_idx in range(len(tbl_index_list)):
        # value of the block_index
        block_index = tbl_index_list[place_idx]
        # Go through the other list to see if you can append the value
        for mini_list in title_list:
            if mini_list[0] == block_index:
                mini_list.append(place_idx)

    # If true return a dictionary of the dataframes
    if pandas_tables:
        df_dict = {}
        # Go through the doc table list and print only the tables that are in the index of the last item of the list
        for table_idx in range(len(doc.tables)):
            # Loop through the title_list to see if the table is to be exported
            for title_idx in title_list:
                # If the table is a match print it out
                if title_idx[2] == table_idx:
                    tbl = doc.tables[table_idx]
                    # Turn the table into a dataframe
                    df = save_table_info(tbl)
                    # Set the df to be the value of the dict and the name to be the appendix name
                    df_dict[title_idx[1]] = df
        # Return the dictionary
        return df_dict

    # Export the tables to an output location of the user's choosing
    else:
        # Go through the doc table list and print only the tables that are in the index of the last item of the list
        for table_idx in range(len(doc.tables)):
            # Loop through the title_list to see if the table is to be exported
            for title_idx in title_list:
                # If the table is a match print it out
                if title_idx[2] == table_idx:
                    # Pulls the table at the appropriate index
                    tbl = doc.tables[table_idx]
                    csv_filename = f'{title_idx[1]}.csv'
                    # Combine the filename with the out path set by the user
                    csv_full_path = path.join(output_location, csv_filename)
                    # Export as a csv
                    save_table_as_csv(tbl, csv_full_path)


# ======================================================================================================================
# Test the function
# ======================================================================================================================
if __name__ == '__main__':
    # Set the variables
    word_path = r'C:\Users\rossc\Documents\ADD_50504_T1DMND0111290_Transportation_3_FCs_W_SAP.docx'
    return_tables = False
    out_file_location = r'C:\Users\rossc\Documents'

    # Test the function
    if return_tables:
        test_dictionary = extract_tables(word_path, return_tables, out_file_location)
        for k, v in test_dictionary.items():
            print(f'Key: {k}, \nValue: \n{v.head()}')
            print('\n\n\n')

    else:
        extract_tables(word_path, return_tables, out_file_location)
